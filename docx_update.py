from docx import Document

# Load the .docx file
doc = Document('daily understandings.docx')

# Read content from the file
for paragraph in doc.paragraphs:
    print(paragraph.text)

# Modify content (optional)
doc.add_paragraph('This is a new paragraph.')

# Save changes
doc.save('your_updated_file.docx')
